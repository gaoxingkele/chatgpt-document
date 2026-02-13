# -*- coding: utf-8 -*-
"""将 Markdown 报告转换为 Word 格式。"""
import re
import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from docx import Document
from docx.shared import Pt


def md_to_docx(md_text: str, docx_path: Path) -> None:
    """将 Markdown 转为 Word：标题层级字号/粗体，段落/列表/表格。"""
    doc = Document()
    doc.styles["Normal"].font.size = Pt(12)
    doc.styles["Normal"].font.name = "宋体"

    lines = md_text.replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 一级标题 # -> 18pt 加粗
        if stripped.startswith("# ") and not stripped.startswith("## "):
            p = doc.add_paragraph(stripped[2:].strip())
            p.style = "Heading 1"
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(18)
            p.runs[0].font.name = "黑体"
            i += 1
            continue

        # 二级标题 ## -> 16pt 加粗
        if stripped.startswith("## ") and not stripped.startswith("### "):
            p = doc.add_paragraph(stripped[3:].strip())
            p.style = "Heading 2"
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(16)
            p.runs[0].font.name = "黑体"
            i += 1
            continue

        # 三级标题 ### -> 14pt 加粗
        if stripped.startswith("### "):
            p = doc.add_paragraph(stripped[4:].strip())
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.name = "黑体"
            i += 1
            continue

        # Markdown 表格：| a | b |
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_md_table(doc, table_lines)
            continue

        # 无序列表 - 或 *
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(text).font.size = Pt(12)
            i += 1
            continue

        # 有序列表 1. 2.
        if re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            p = doc.add_paragraph(style="List Number")
            p.add_run(text).font.size = Pt(12)
            i += 1
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24)
        p.add_run(stripped).font.size = Pt(12)
        i += 1

    doc.save(str(docx_path))


def _add_md_table(doc, table_lines: list) -> None:
    """解析 Markdown 表格行并写入 docx。"""
    if len(table_lines) < 2:
        return
    rows = []
    for ln in table_lines:
        cells = [c.strip() for c in ln.split("|") if c.strip()]
        if cells and not all(re.match(r"^[-:]+$", c) for c in cells):
            rows.append(cells)
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for ri, row_cells in enumerate(rows):
        for ci, cell_text in enumerate(row_cells):
            if ci < col_count:
                cell = table.rows[ri].cells[ci]
                cell.text = cell_text
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10.5)
                if ri == 0:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.bold = True
    doc.add_paragraph()


if __name__ == "__main__":
    md_path = PROJECT_ROOT / "output" / "reports" / "web3_integral_report_v2.md"
    docx_path = PROJECT_ROOT / "output" / "reports" / "web3_integral_report_v2.docx"
    if md_path.is_file():
        text = md_path.read_text(encoding="utf-8", errors="replace")
        md_to_docx(text, docx_path)
        print(f"[完成] Word 文档已保存: {docx_path}")
    else:
        print(f"[错误] 找不到文件: {md_path}")
