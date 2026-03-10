# -*- coding: utf-8 -*-
"""Markdown → Word 转换工具（增强版）。

支持：标题、段落、加粗、斜体、超链接、引用块、嵌套列表、
水平线、代码块、图片、表格列对齐、目录、页眉页脚、封面页、脚注、
数学公式（$inline$ 和 $$display$$）。
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


# --------------- 数学公式 ---------------

def _add_inline_math(paragraph, latex: str):
    """在段落中插入行内数学公式（OMML 格式）。"""
    try:
        from src.utils.omml_converter import latex_to_omml
        from lxml import etree
        omath = latex_to_omml(latex)
        # 设置 OMML 命名空间声明
        OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
        omath_str = etree.tostring(omath, encoding="unicode")
        # 添加命名空间前缀
        if f'xmlns:m="{OMML_NS}"' not in omath_str:
            omath_str = omath_str.replace(
                f"{{{OMML_NS}}}",
                "m:",
            )
            omath_str = omath_str.replace(
                "<m:oMath",
                f'<m:oMath xmlns:m="{OMML_NS}"',
                1,
            )
        omath_el = parse_xml(omath_str)
        paragraph._element.append(omath_el)
    except Exception:
        # 回退：作为斜体文本
        r = paragraph.add_run(latex)
        r.italic = True
        r.font.size = Pt(12)


def _add_display_math(doc, latex: str):
    """添加块级数学公式（独立居中段落）。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_inline_math(p, latex)


def _is_display_math_block(lines: list, i: int) -> tuple:
    """
    检测 $$ ... $$ 块级公式。
    返回 (latex_content, end_line_index) 或 (None, i)。
    """
    stripped = lines[i].strip()
    # 单行 $$ ... $$
    m = re.match(r'^\$\$(.+)\$\$$', stripped)
    if m:
        return m.group(1).strip(), i + 1

    # 多行 $$ 开始
    if stripped == "$$":
        j = i + 1
        parts = []
        while j < len(lines):
            if lines[j].strip() == "$$":
                return "\n".join(parts).strip(), j + 1
            parts.append(lines[j])
            j += 1
        # 未闭合的 $$：作为文本
        return None, i

    return None, i


# --------------- 内联格式解析 ---------------

def _add_hyperlink(paragraph, url: str, text: str):
    """在段落中添加 Word 超链接。"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = parse_xml(
        f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'<w:r><w:rPr><w:rStyle w:val="Hyperlink"/>'
        f'<w:color w:val="0563C1"/><w:u w:val="single"/>'
        f'</w:rPr><w:t>{_xml_escape(text)}</w:t></w:r></w:hyperlink>'
    )
    paragraph._element.append(hyperlink)


def _xml_escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")


def _add_runs_with_formatting(paragraph, text: str, font_size, font_name: str = "宋体") -> None:
    """解析文本中的 $公式$、**加粗**、*斜体*、[链接](url)，并添加到段落。"""
    # 先拆分行内公式 $...$ （不匹配 $$），再处理其他格式
    # 注意：不贪婪匹配，且 $ 不能紧跟 $ 或前接 $
    math_pattern = re.compile(r'(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)')
    parts = []
    last = 0
    for m in math_pattern.finditer(text):
        if m.start() > last:
            parts.append(("text", text[last:m.start()]))
        parts.append(("math", m.group(1)))
        last = m.end()
    if last < len(text):
        parts.append(("text", text[last:]))

    for ptype, pval in parts:
        if ptype == "math":
            _add_inline_math(paragraph, pval)
        else:
            _add_formatted_text_segment(paragraph, pval, font_size, font_name)


def _add_formatted_text_segment(paragraph, text: str, font_size, font_name: str = "宋体") -> None:
    """解析文本片段中的 **加粗**、*斜体*、[链接](url)。"""
    pattern = re.compile(r'(\[([^\]]+)\]\(([^)]+)\)|\*\*([^*]+)\*\*|\*([^*]+)\*)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            r.font.size = font_size
            r.font.name = font_name
        if m.group(2) and m.group(3):
            _add_hyperlink(paragraph, m.group(3), m.group(2))
        elif m.group(4):
            r = paragraph.add_run(m.group(4))
            r.bold = True
            r.font.size = font_size
            r.font.name = font_name
        elif m.group(5):
            r = paragraph.add_run(m.group(5))
            r.italic = True
            r.font.size = font_size
            r.font.name = font_name
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        r.font.size = font_size
        r.font.name = font_name


def _add_runs_with_bold(paragraph, text: str, font_size) -> None:
    """向后兼容：仅解析 **加粗**。"""
    _add_runs_with_formatting(paragraph, text, font_size)


# --------------- 表格增强 ---------------

def _parse_table_alignment(separator_line: str) -> list[str]:
    """从分隔行解析列对齐方式：left / center / right。"""
    cells = [c.strip() for c in separator_line.split("|") if c.strip()]
    alignments = []
    for c in cells:
        if not re.match(r'^[-:]+$', c):
            alignments.append("left")
            continue
        if c.startswith(":") and c.endswith(":"):
            alignments.append("center")
        elif c.endswith(":"):
            alignments.append("right")
        else:
            alignments.append("left")
    return alignments


def _add_md_table(doc, table_lines: list) -> None:
    """解析 Markdown 表格行并写入 docx（支持列对齐）。"""
    if len(table_lines) < 2:
        return

    # 检测分隔行并获取对齐方式
    alignments = []
    rows = []
    for ln in table_lines:
        cells = [c.strip() for c in ln.split("|") if c.strip()]
        if cells and all(re.match(r'^[-:]+$', c) for c in cells):
            alignments = _parse_table_alignment(ln)
            continue
        if cells:
            rows.append(cells)
    if not rows:
        return

    col_count = max(len(r) for r in rows)
    # 补齐对齐列表
    while len(alignments) < col_count:
        alignments.append("left")

    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    align_map = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}

    for ri, row_cells in enumerate(rows):
        for ci, cell_text in enumerate(row_cells):
            if ci < col_count:
                cell = table.rows[ri].cells[ci]
                cell.text = cell_text
                for p in cell.paragraphs:
                    if ci < len(alignments):
                        p.alignment = align_map.get(alignments[ci], WD_ALIGN_PARAGRAPH.LEFT)
                    for r in p.runs:
                        r.font.size = Pt(10.5)
                if ri == 0:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.bold = True
    doc.add_paragraph()


# --------------- 代码块 ---------------

def _add_code_block(doc, code_lines: list[str]) -> None:
    """添加代码块：Courier New 10pt，浅灰背景。"""
    code_text = "\n".join(code_lines)
    p = doc.add_paragraph()
    # 浅灰背景
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>')
    p._element.get_or_add_pPr().append(shading)
    r = p.add_run(code_text)
    r.font.name = "Courier New"
    r.font.size = Pt(10)


# --------------- 图片 ---------------

def _add_image(doc, alt_text: str, img_path: str) -> None:
    """添加图片，路径不存在时插入占位文字。"""
    resolved = Path(img_path)
    if not resolved.is_absolute():
        resolved = Path.cwd() / img_path
    if resolved.is_file():
        try:
            doc.add_picture(str(resolved), width=Inches(5.5))
            if alt_text:
                p = doc.add_paragraph(alt_text)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.italic = True
            return
        except Exception:
            pass
    # 占位
    p = doc.add_paragraph(f"[图片: {alt_text or img_path}]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.size = Pt(10)
        r.italic = True


# --------------- 目录 ---------------

def _add_toc(doc) -> None:
    """在文档开头插入 Word 目录域代码（打开文件后按 F9 更新）。"""
    p = doc.add_paragraph()
    r = p.add_run("目 录")
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = "黑体"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 插入 TOC 域
    fld_xml = (
        f'<w:p {nsdecls("w")}>'
        '<w:r><w:fldChar w:fldCharType="begin"/></w:r>'
        '<w:r><w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText></w:r>'
        '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
        '<w:r><w:t>[请右键目录区域，选择"更新域"以生成目录]</w:t></w:r>'
        '<w:r><w:fldChar w:fldCharType="end"/></w:r>'
        '</w:p>'
    )
    doc._element.body.append(parse_xml(fld_xml))
    # 分页
    p2 = doc.add_paragraph()
    r2 = p2.add_run()
    r2._element.append(parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>'))


# --------------- 页眉页脚 ---------------

def _add_header_footer(doc, title: str = "") -> None:
    """添加页眉（报告标题）和页脚（页码）。"""
    section = doc.sections[0]

    # 页眉
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = title
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in hp.runs:
        r.font.size = Pt(9)
        r.font.name = "宋体"

    # 页脚：页码
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 添加 PAGE 域代码
    fld_begin = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:sz w:val="18"/></w:rPr><w:fldChar w:fldCharType="begin"/></w:r>')
    fld_instr = parse_xml(f'<w:r {nsdecls("w")}><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>')
    fld_end = parse_xml(f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="end"/></w:r>')
    fp._element.append(fld_begin)
    fp._element.append(fld_instr)
    fp._element.append(fld_end)


# --------------- 封面页 ---------------

def _add_cover_page(doc, title: str, keywords: list[str] = None) -> None:
    """添加封面页：居中标题 + 日期 + 关键词。"""
    import datetime

    # 空行间距
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(26)
    r.font.name = "黑体"

    doc.add_paragraph()

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(datetime.date.today().strftime("%Y 年 %m 月 %d 日"))
    r2.font.size = Pt(14)
    r2.font.name = "宋体"

    if keywords:
        doc.add_paragraph()
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run("关键词：" + "、".join(keywords))
        r3.font.size = Pt(12)
        r3.font.name = "宋体"

    # 分页
    p4 = doc.add_paragraph()
    r4 = p4.add_run()
    r4._element.append(parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>'))


# --------------- 脚注处理 ---------------

_FOOTNOTE_PATTERN = re.compile(r'\[\^(\d+)\]')
_FOOTNOTE_DEF_PATTERN = re.compile(r'^\[\^(\d+)\]:\s*(.+)$', re.MULTILINE)


def _collect_footnote_defs(md_text: str) -> dict[str, str]:
    """收集脚注定义 [^N]: text。"""
    defs = {}
    for m in _FOOTNOTE_DEF_PATTERN.finditer(md_text):
        defs[m.group(1)] = m.group(2).strip()
    return defs


# --------------- 水平线检测 ---------------

def _is_horizontal_rule(stripped: str) -> bool:
    """检测 Markdown 水平线：---、***、___（至少3个连续字符）。"""
    return bool(re.match(r'^[-*_]{3,}\s*$', stripped))


# --------------- 主转换函数 ---------------

def _get_fonts():
    """根据报告语言返回字体配置。"""
    try:
        from config import REPORT_LANGUAGE
        lang = REPORT_LANGUAGE
    except (ImportError, AttributeError):
        lang = "zh"
    if lang == "en":
        return {"body": "Times New Roman", "heading": "Arial"}
    return {"body": "宋体", "heading": "黑体"}


def md_to_docx(md_text: str, docx_path: Path) -> None:
    """将 Markdown 转为 Word（增强版）。

    支持：标题层级、段落、加粗/斜体、超链接、引用块、嵌套列表、
    水平线、代码块、图片、表格列对齐、目录、页眉页脚、封面页。
    """
    fonts = _get_fonts()
    body_font = fonts["body"]
    heading_font = fonts["heading"]

    doc = Document()
    doc.styles["Normal"].font.size = Pt(12)
    doc.styles["Normal"].font.name = body_font

    # 收集脚注定义
    footnote_defs = _collect_footnote_defs(md_text)

    # 提取标题和关键词用于封面/页眉
    title_match = re.search(r'^#\s+(.+)$', md_text, re.MULTILINE)
    doc_title = title_match.group(1).strip() if title_match else ""
    kw_match = re.search(r'\*\*关键词\*\*[：:]\s*(.+)$', md_text, re.MULTILINE)
    keywords = [k.strip() for k in kw_match.group(1).split("、")] if kw_match else []

    # 封面页
    if doc_title:
        _add_cover_page(doc, doc_title, keywords)

    # 目录
    _add_toc(doc)

    # 页眉页脚
    _add_header_footer(doc, doc_title)

    lines = md_text.replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 跳过脚注定义行
        if _FOOTNOTE_DEF_PATTERN.match(stripped):
            i += 1
            continue

        # 一级标题 # -> 18pt 加粗
        if stripped.startswith("# ") and not stripped.startswith("## "):
            p = doc.add_paragraph(stripped[2:].strip())
            p.style = "Heading 1"
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(18)
            p.runs[0].font.name = heading_font
            i += 1
            continue

        # 二级标题 ## -> 16pt 加粗
        if stripped.startswith("## ") and not stripped.startswith("### "):
            p = doc.add_paragraph(stripped[3:].strip())
            p.style = "Heading 2"
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(16)
            p.runs[0].font.name = heading_font
            i += 1
            continue

        # 三级标题 ### -> 14pt 加粗
        if stripped.startswith("### "):
            p = doc.add_paragraph(stripped[4:].strip())
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.name = heading_font
            i += 1
            continue

        # 块级数学公式 $$ ... $$
        if stripped.startswith("$$"):
            latex, next_i = _is_display_math_block(lines, i)
            if latex is not None:
                _add_display_math(doc, latex)
                i = next_i
                continue

        # 代码块 ```
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # 跳过结束 ```
            _add_code_block(doc, code_lines)
            continue

        # 水平线 --- / *** / ___
        if _is_horizontal_rule(stripped):
            p = doc.add_paragraph()
            # 底部边框
            pPr = p._element.get_or_add_pPr()
            border_xml = (
                f'<w:pBdr {nsdecls("w")}>'
                '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="999999"/>'
                '</w:pBdr>'
            )
            pPr.append(parse_xml(border_xml))
            i += 1
            continue

        # 图片 ![alt](path)
        img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$', stripped)
        if img_match:
            _add_image(doc, img_match.group(1), img_match.group(2))
            i += 1
            continue

        # Markdown 表格：| a | b |
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_md_table(doc, table_lines)
            continue

        # 引用块 > text
        if stripped.startswith("> ") or stripped == ">":
            text = stripped[2:].strip() if len(stripped) > 2 else ""
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.27)  # 36pt ≈ 1.27cm
            _add_runs_with_formatting(p, text, Pt(12))
            for r in p.runs:
                r.italic = True
            i += 1
            continue

        # 嵌套无序列表（2+ 空格缩进）
        indent_match = re.match(r'^(\s{2,})[-*]\s+(.+)$', line)
        if indent_match:
            indent_level = len(indent_match.group(1)) // 2
            text = indent_match.group(2).strip()
            style = "List Bullet 2" if indent_level >= 1 else "List Bullet"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Cm(1.27 * indent_level)
            _add_runs_with_formatting(p, text, Pt(12))
            i += 1
            continue

        # 无序列表 - 或 *
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_formatting(p, text, Pt(12))
            i += 1
            continue

        # 有序列表 1. 2.
        if re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _add_runs_with_formatting(p, text, Pt(12))
            i += 1
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 普通段落（支持内联格式）
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24)
        _add_runs_with_formatting(p, stripped, Pt(12))
        i += 1

    # 添加脚注说明（若有定义）
    if footnote_defs:
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run("注释")
        r.bold = True
        r.font.size = Pt(14)
        r.font.name = "黑体"
        for num, text in sorted(footnote_defs.items(), key=lambda x: int(x[0])):
            fp = doc.add_paragraph()
            fp.paragraph_format.left_indent = Cm(0.63)
            r = fp.add_run(f"[{num}] {text}")
            r.font.size = Pt(10)
            r.font.name = "宋体"

    doc.save(str(docx_path))


def save_docx_safe(md_text: str, docx_path: Path) -> Path:
    """md_to_docx 的安全包装：PermissionError 时自动追加 _new 后缀并记录日志。返回实际保存路径。"""
    from src.utils.log import log as _log
    try:
        md_to_docx(md_text, docx_path)
        return docx_path
    except PermissionError:
        stem = docx_path.stem
        alt_path = docx_path.with_name(f"{stem}_new.docx")
        md_to_docx(md_text, alt_path)
        _log(f"[提示] 原文件可能被占用，已保存为: {alt_path.name}")
        return alt_path
