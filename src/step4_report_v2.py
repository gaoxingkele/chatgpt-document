# -*- coding: utf-8 -*-
"""
Step4: 根据专家意见与原始语料，对报告 1.0 整改生成报告 2.0；
要求：保留 ChatGPT 论述逻辑、5~7 章、去重简练、专业严谨；输出 Markdown 与 Word。
"""
import sys
import time
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent


def _log(msg: str):
    ts = time.strftime("%H:%M:%S", time.localtime())
    print(f"[{ts}] {msg}", flush=True)
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from config import REPORT_DIR, EXPERT_DIR, RAW_DIR
from src.kimi_client import chat
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


def _load_expert_combined(base: str) -> str:
    p = EXPERT_DIR / f"{base}_专家意见汇总.md"
    if not p.is_file():
        raise FileNotFoundError(f"专家意见汇总不存在: {p}，请先运行 Step3")
    return p.read_text(encoding="utf-8", errors="replace")


def _load_hallucination_list(base: str) -> str:
    """加载幻觉清单（若存在）。"""
    p = EXPERT_DIR / f"{base}_专家4_幻觉清单.md"
    if p.is_file():
        return p.read_text(encoding="utf-8", errors="replace")
    return ""


def _load_raw_content(raw_path: Path, max_chars: int = 120000) -> str:
    """加载原始语料，用于补充论述逻辑与篇幅约束。"""
    if not raw_path or not Path(raw_path).is_file():
        return ""
    text = Path(raw_path).read_text(encoding="utf-8", errors="replace")
    return text[:max_chars] + ("\n\n[已截断]" if len(text) > max_chars else "")


def _parse_report_v1_chapters(text: str) -> tuple[str, list[tuple[str, str]]]:
    """
    解析报告 1.0：提取正文前的头部（标题、摘要、关键词），及章节列表 [(章标题, 章正文), ...]。
    章标题匹配 ## 一、 ## 二、 ... 或 ## 1. ## 2. ...
    """
    # 匹配 ## 一、xxx 或 ## 1. xxx
    pattern = re.compile(r"^##\s+[一二三四五六七八九十]+、.+$|^##\s+\d+\.\s+.+$", re.MULTILINE)
    matches = list(pattern.finditer(text))
    if not matches:
        return text, []

    header = text[: matches[0].start()].strip()
    chapters: list[tuple[str, str]] = []
    for i, m in enumerate(matches):
        title = m.group(0).strip()
        body_start = m.end()
        body_end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[body_start:body_end].strip()
        chapters.append((title, body))
    return header, chapters


def _api_revise_chapter(
    chapter_title: str,
    chapter_body: str,
    expert_text: str,
    hallucination_text: str,
    raw_chunk: str,
    target_chars: int,
    chapter_idx: int,
    total_chapters: int,
) -> str:
    """对单章进行整改，返回该章完整正文（含章标题）。强调篇幅必须达标。"""
    hall_section = ""
    if hallucination_text:
        hall_section = f"""
【幻觉清单】必须删除以下内容，不得出现在本章：
{hallucination_text[:8000]}
"""
    raw_section = ""
    if raw_chunk:
        raw_section = f"""
【原始语料】（供参考，优先保留论证、案例、数据）
---
{raw_chunk[:35000]}
---
"""
    prompt = f"""请对《深度调查报告 1.0》的**第 {chapter_idx}/{total_chapters} 章**进行整改，输出该章的完整正文。

【本章标题】{chapter_title}

【本章正文（报告 1.0）】
---
{chapter_body}
---

【专家评审意见】（采纳可执行的改进）
---
{expert_text[:25000]}
---
{hall_section}{raw_section}

【极其重要的篇幅要求（必须遵守）】
- 本章输出字数**不少于 {target_chars} 字**。禁止压缩、禁止将多段合并成一句或要点罗列。
- 重写、去重、理顺逻辑，但**不要删减论证、案例、表格、数据**。
- 直接输出本章完整正文，以 `## {chapter_title}` 开头，使用 Markdown（### 等）。不要 JSON 或多余说明。"""

    resp = chat(
        [
            {
                "role": "system",
                "content": "你是专业的研究报告修订专家。核心原则：1) 篇幅必须充足，每章不少于目标字数；2) 保留论述逻辑与案例丰富度；3) 重写而非压缩；4) 吸收专家意见。输出严格为 Markdown。",
            },
            {"role": "user", "content": prompt},
        ],
        max_tokens=16384,
        temperature=0.4,
    )
    return resp.strip()


def run_report_v2_and_docx(
    report_v1_path: Path,
    expert_combined_path: Path = None,
    output_basename: str = None,
    raw_path: Path = None,
) -> dict:
    """
    根据报告 1.0、专家意见汇总与（可选）原始语料，生成报告 2.0 的 Markdown，再转为 Word。
    返回 report_v2_path（.md）与 docx_path（.docx）。
    """
    report_v1_path = Path(report_v1_path)
    if not report_v1_path.is_file():
        raise FileNotFoundError(f"报告 1.0 不存在: {report_v1_path}")
    report_v1_text = report_v1_path.read_text(encoding="utf-8", errors="replace")
    base = output_basename or report_v1_path.stem.replace("_report_v1", "")

    if expert_combined_path and Path(expert_combined_path).is_file():
        expert_text = Path(expert_combined_path).read_text(encoding="utf-8", errors="replace")
    else:
        expert_text = _load_expert_combined(base)

    hallucination_text = _load_hallucination_list(base)
    if hallucination_text:
        _log(f"已加载幻觉清单，共约 {len(hallucination_text)} 字")

    raw_text = _load_raw_content(raw_path, 120000) if raw_path else ""
    raw_full_len = len(Path(raw_path).read_text(encoding="utf-8", errors="replace")) if raw_path and Path(raw_path).is_file() else len(raw_text)
    target_min_chars = max(16000, int(raw_full_len * 0.6))

    header, chapters = _parse_report_v1_chapters(report_v1_text)
    num_chapters = len(chapters)
    target_per_chapter = max(2000, target_min_chars // num_chapters) if num_chapters else target_min_chars

    _log("=" * 60)
    _log("Step4 报告 2.0：开始（分章整改模式）")
    _log(f"报告 1.0: 约 {len(report_v1_text)} 字 | 原始语料: 约 {raw_full_len} 字 | 字数目标: ≥{target_min_chars} 字")
    _log(f"章节数: {num_chapters} | 每章目标: ≥{target_per_chapter} 字")
    _log("=" * 60)
    t0 = time.time()

    revised_parts: list[str] = []
    raw_len = len(raw_text)
    for idx, (ch_title, ch_body) in enumerate(chapters):
        _log(f"--- 整改第 {idx + 1}/{num_chapters} 章: {ch_title[:40]}...")
        # 原始语料按章均分，便于每章获取相关上下文
        start_pos = idx * raw_len // num_chapters if raw_len else 0
        end_pos = (idx + 1) * raw_len // num_chapters if raw_len else raw_len
        raw_chunk = raw_text[start_pos:end_pos] if raw_text else ""
        revised = _api_revise_chapter(
            ch_title,
            ch_body[:15000],
            expert_text,
            hallucination_text,
            raw_chunk,
            target_per_chapter,
            idx + 1,
            num_chapters,
        )
        _log(f"    完成，输出约 {len(revised)} 字")
        revised_parts.append(revised)

    # 拼接：头部 + 各章
    report_v2_body = "\n\n".join(revised_parts)
    report_v2_text = f"{header}\n\n{report_v2_body}".strip()

    _log(f"分章整改完成，总耗时 {time.time()-t0:.1f}s，报告 2.0 约 {len(report_v2_text)} 字")

    report_v2_path = REPORT_DIR / f"{base}_report_v2.md"
    report_v2_path.write_text(report_v2_text, encoding="utf-8")
    _log(f"报告 2.0 (Markdown) 已保存: {report_v2_path.name}")

    # 转为 Word：标题层级字号与粗体，正文段落与列表
    _log("导出 Word：报告 2.0 → .docx")
    docx_path = REPORT_DIR / f"{base}_report_v2.docx"
    try:
        md_to_docx(report_v2_text, docx_path)
    except PermissionError:
        alt_path = REPORT_DIR / f"{base}_report_v2_new.docx"
        md_to_docx(report_v2_text, alt_path)
        docx_path = alt_path
        _log(f"[提示] 原文件可能被占用，已保存为: {docx_path.name}")
    _log(f"Step4 完成：报告 2.0 (Word) 已保存 {docx_path.name}")

    return {
        "report_v2_path": str(report_v2_path),
        "docx_path": str(docx_path),
        "report_v2_text": report_v2_text,
    }


def _add_runs_with_bold(paragraph, text: str, font_size) -> None:
    """将文本中的 **xxx** 解析为加粗，其余为普通。"""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
            r.font.size = font_size
            r.font.name = "宋体"
        elif part:
            r = paragraph.add_run(part)
            r.font.size = font_size
            r.font.name = "宋体"


def md_to_docx(md_text: str, docx_path: Path) -> None:
    """将 Markdown 转为 Word：标题层级字号/粗体，段落/列表/表格，正文内 **加粗**。"""
    doc = Document()
    doc.styles["Normal"].font.size = Pt(12)
    doc.styles["Normal"].font.name = "宋体"

    lines = md_text.replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 一级标题 # -> 18pt 加粗
        if stripped.startswith("# ") and not stripped.startswith("## "):
            p = doc.add_paragraph(stripped[2:].strip())
            p.style = "Heading 1"
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(18)
            p.runs[0].font.name = "黑体"
            i += 1
            continue

        # 二级标题 ## -> 16pt 加粗
        if stripped.startswith("## ") and not stripped.startswith("### "):
            p = doc.add_paragraph(stripped[3:].strip())
            p.style = "Heading 2"
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(16)
            p.runs[0].font.name = "黑体"
            i += 1
            continue

        # 三级标题 ### -> 14pt 加粗
        if stripped.startswith("### "):
            p = doc.add_paragraph(stripped[4:].strip())
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.name = "黑体"
            i += 1
            continue

        # Markdown 表格：| a | b |
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_md_table(doc, table_lines)
            continue

        # 无序列表 - 或 *
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_bold(p, text, Pt(12))
            i += 1
            continue

        # 有序列表 1. 2.
        if re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _add_runs_with_bold(p, text, Pt(12))
            i += 1
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 普通段落（支持 **加粗**）
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24)
        _add_runs_with_bold(p, stripped, Pt(12))
        i += 1

    doc.save(str(docx_path))


def _add_md_table(doc, table_lines: list) -> None:
    """解析 Markdown 表格行并写入 docx。"""
    if len(table_lines) < 2:
        return
    rows = []
    for ln in table_lines:
        cells = [c.strip() for c in ln.split("|") if c.strip()]
        if cells and not all(re.match(r"^[-:]+$", c) for c in cells):
            rows.append(cells)
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for ri, row_cells in enumerate(rows):
        for ci, cell_text in enumerate(row_cells):
            if ci < col_count:
                cell = table.rows[ri].cells[ci]
                cell.text = cell_text
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10.5)
                if ri == 0:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.bold = True
    doc.add_paragraph()


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="根据专家意见生成报告2.0并导出Word")
    parser.add_argument("report_v1", type=Path, help="报告 1.0 路径")
    parser.add_argument("-e", "--expert-file", type=Path, default=None, help="专家意见汇总路径（可选）")
    parser.add_argument("-r", "--raw-file", type=Path, default=None, help="原始语料路径（可选，用于补充论述逻辑）")
    parser.add_argument("-o", "--output-base", default=None, help="输出文件名前缀")
    args = parser.parse_args()
    raw = Path(args.raw_file) if args.raw_file else None
    if raw and not raw.is_absolute():
        raw = RAW_DIR / raw.name
    run_report_v2_and_docx(args.report_v1, args.expert_file, args.output_base, raw)
