# -*- coding: utf-8 -*-
"""
语料提取：Word、PDF、图片等格式。
Word/PDF 本地提取文本和嵌入图片；图片提交云端 Vision API 处理。
"""
import base64
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable, List, Optional


@dataclass
class ExtractionResult:
    """富提取结果：文本 + 图片列表。"""
    text: str = ""
    images: List[dict] = field(default_factory=list)


def _get_vision_api() -> Callable[[list], str]:
    from src.llm_client import chat_vision
    return chat_vision


def _guess_image_ext(content_type: str, blob: bytes) -> str:
    """根据 content type 或魔数推断图片扩展名。"""
    ct = content_type.lower() if content_type else ""
    if "png" in ct:
        return ".png"
    if "jpeg" in ct or "jpg" in ct:
        return ".jpg"
    if "gif" in ct:
        return ".gif"
    if "webp" in ct:
        return ".webp"
    if "bmp" in ct:
        return ".bmp"
    if "tiff" in ct or "tif" in ct:
        return ".tiff"
    if "emf" in ct or "x-emf" in ct:
        return ".emf"
    if "wmf" in ct or "x-wmf" in ct:
        return ".wmf"
    # 魔数检测
    if blob[:8] == b'\x89PNG\r\n\x1a\n':
        return ".png"
    if blob[:2] == b'\xff\xd8':
        return ".jpg"
    if blob[:4] == b'GIF8':
        return ".gif"
    if blob[:4] == b'RIFF' and blob[8:12] == b'WEBP':
        return ".webp"
    return ".png"


def extract_from_docx_rich(file_path: Path) -> ExtractionResult:
    """
    从 Word (.docx) 提取文本和嵌入图片。
    图片位置用 ![image](assets/img_NNN.ext) 占位，图片数据在 images 列表中。
    """
    from docx import Document
    from lxml import etree

    doc = Document(file_path)

    # XML 命名空间
    nsmap = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    }
    OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

    # 收集所有图片关系 (rId → image blob + content_type)
    image_rels = {}
    try:
        for rel_id, rel in doc.part.rels.items():
            if "image" in (rel.reltype or ""):
                try:
                    blob = rel.target_part.blob
                    ct = rel.target_part.content_type or ""
                    image_rels[rel_id] = (blob, ct)
                except Exception:
                    pass
    except Exception:
        pass

    img_counter = 0
    images = []
    parts = []

    for para in doc.paragraphs:
        para_parts = []
        xml = para._element

        # 遍历段落子元素
        for child in xml:
            tag = etree.QName(child.tag).localname if isinstance(child.tag, str) else ""

            if tag == "r":
                # 文本 run：检查是否含图片
                drawings = child.findall(".//wp:inline", nsmap) + child.findall(".//wp:anchor", nsmap)
                if drawings:
                    for drawing in drawings:
                        blip = drawing.find(".//a:blip", nsmap)
                        if blip is not None:
                            r_embed = blip.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                            )
                            if r_embed and r_embed in image_rels:
                                blob, ct = image_rels[r_embed]
                                img_counter += 1
                                ext = _guess_image_ext(ct, blob)
                                filename = f"img_{img_counter:03d}{ext}"
                                images.append({"data": blob, "filename": filename})
                                para_parts.append(f"![image](assets/{filename})")

                # 文本内容
                text_nodes = child.findall(".//w:t", nsmap)
                for t in text_nodes:
                    if t.text:
                        para_parts.append(t.text)

            elif tag == "drawing":
                # 顶层 drawing 元素
                blip = child.find(".//a:blip", nsmap)
                if blip is not None:
                    r_embed = blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )
                    if r_embed and r_embed in image_rels:
                        blob, ct = image_rels[r_embed]
                        img_counter += 1
                        ext = _guess_image_ext(ct, blob)
                        filename = f"img_{img_counter:03d}{ext}"
                        images.append({"data": blob, "filename": filename})
                        para_parts.append(f"![image](assets/{filename})")

            elif tag == "oMath":
                # 行内公式
                try:
                    from src.utils.omml_converter import omml_to_latex
                    latex = omml_to_latex(child)
                    if latex:
                        para_parts.append(f"${latex}$")
                except Exception:
                    pass

            elif tag == "oMathPara":
                # 块级公式
                try:
                    from src.utils.omml_converter import omml_to_latex
                    omath = child.find(f"{{{OMML_NS}}}oMath")
                    if omath is not None:
                        latex = omml_to_latex(omath)
                        if latex:
                            para_parts.append(f"\n$$\n{latex}\n$$\n")
                except Exception:
                    pass

        text = "".join(para_parts).strip()
        if text:
            parts.append(text)

    # 表格
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return ExtractionResult(text="\n\n".join(parts), images=images)


def extract_from_docx(file_path: Path) -> str:
    """从 Word (.docx) 提取文本（向后兼容）。"""
    return extract_from_docx_rich(file_path).text


def extract_from_pdf_rich(file_path: Path) -> ExtractionResult:
    """
    从 PDF 提取文本和嵌入图片。
    图片附在对应页面文本之后，用 ![image](assets/img_NNN.ext) 占位。
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise ImportError("请安装 pypdf: pip install pypdf")

    reader = PdfReader(file_path)
    parts = []
    images = []
    img_counter = 0

    for page in reader.pages:
        # 提取文本
        text = page.extract_text()
        page_text = text.strip() if text else ""

        # 提取图片
        page_img_refs = []
        try:
            for image in page.images:
                img_counter += 1
                blob = image.data
                ext = _guess_image_ext("", blob)
                # 用原始文件名扩展名（如果有）
                if image.name:
                    orig_ext = Path(image.name).suffix.lower()
                    if orig_ext in (".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff"):
                        ext = orig_ext
                filename = f"img_{img_counter:03d}{ext}"
                images.append({"data": blob, "filename": filename})
                page_img_refs.append(f"![image](assets/{filename})")
        except Exception:
            pass

        if page_text or page_img_refs:
            segment = page_text
            if page_img_refs:
                segment += "\n\n" + "\n\n".join(page_img_refs) if segment else "\n\n".join(page_img_refs)
            parts.append(segment)

    return ExtractionResult(text="\n\n".join(parts), images=images)


def extract_from_pdf(file_path: Path) -> str:
    """从 PDF 提取文本（向后兼容）。"""
    return extract_from_pdf_rich(file_path).text


def extract_from_image(file_path: Path, vision_api: Optional[Callable[[list], str]] = None) -> str:
    """
    从图片提取内容：编码为 base64 后提交云端 Vision API 处理。
    vision_api: 接受 messages 列表，返回文本；未传入则使用 kimi_client.chat_vision。
    """
    if vision_api is None:
        vision_api = _get_vision_api()
    path = Path(file_path)
    raw = path.read_bytes()
    b64 = base64.b64encode(raw).decode("ascii")
    suffix = path.suffix.lower()
    mime = {
        ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".png": "image/png", ".gif": "image/gif",
        ".webp": "image/webp", ".bmp": "image/bmp",
    }.get(suffix, "image/jpeg")
    data_url = f"data:{mime};base64,{b64}"

    content = [
        {
            "type": "text",
            "text": "请提取此图片中的全部文字、图表、表格及关键信息，输出为可直接用于语料整理的纯文本。若为截图或文档图片，请完整还原文字内容。",
        },
        {"type": "image_url", "image_url": {"url": data_url}},
    ]
    messages = [
        {"role": "system", "content": "你是专业的内容提取专家。任务：从图片中提取全部可读文字与关键信息，输出为结构清晰的纯文本，便于后续语料整理。"},
        {"role": "user", "content": content},
    ]
    return vision_api(messages)
